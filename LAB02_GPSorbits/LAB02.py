'''
Positioning and Location Based Services
A.A. 2022/2023
Exercise 2:  GPS orbits computation

Marianna Alghisi

References:
    - http://www.navipedia.net/index.php/Clock_Modelling
    - http://www.ajgeomatics.com/SatelliteOrbitsRev0828-2010.pdf
    - http://www.navipedia.net/index.php/GPS_and_Galileo_Satellite_Coordinates_Computation
'''
 
# Import all the required libraries and external script containing the functions
import numpy as np
import matplotlib.pyplot as plt
from PLBS import *

'''pip install docx
from docx import Document
from docx.shared import Inches
from pandas.compat import StringIO'''


# Load Data from Almanac of satellite SVN 63, PRN 01 (Block IIR) for 2016
# set of quasi keplerian parameters to predict the orbit of a GPS satellite
dt0= -7.661711424589e-05
dt1= -3.183231456205e-12
dt2=  0.000000000000e+00

# Keplerian orbit 6 parameters
a= 5.153650835037e+03**2
e= 3.841053112410e-03 
M0= 1.295004883409e+00  
Omega0= -2.241692424630e-01  
i0= 9.634782624741e-01 
w0= 9.419793734505e-01 
wdot= 0.0  
idot= -7.286017777600e-11
Omegadot= -8.386063598924e-09  
OmegaEdot = 7.2921151467e-05  

GMe = 3.986005e+14 # Earth gravitational constant [m^3 sec^-2]

# Create a list containing all the epochs
t_start = 0
t_end = 86340
t_step = 30
t = list(range(t_start, t_end, t_step))

'''
1) Compute clock offsets and plot it
'''
# Loop on t and append the elements to a list
clock_offsets = [ dt0 + dt1*(i - t_start) for i in t]

# plot
fig, ax = plt.subplots(figsize=(10,6))
ax.set(xlabel='Seconds in a day', ylabel='Clock-offset')
ax.plot(t, clock_offsets, '-', color='blue')
 
'''
2) Compute the position of the satellite in Cartesian ITRF [X, Y, Z]
'''
coord_ORS  = np.zeros(shape=(3,len(t))) # Initialize to zero matrices of sat_coord_ORS with dimensions 3-by-n_epochs
coord_ITRF = np.zeros(shape=(3,len(t))) # Initialize to zero matrices of sat_coord_ITRF with dimensions 3-by-n_epochs
r_mean = np.zeros(shape=len(t)) 
n = np.sqrt(GMe/(a**3)) # Mean angular velocity

for Dt in range(0,len(t)):
    Dt_i = t[Dt]
    M = M0 + n*(Dt_i - t_start) 
    eta = eccAnomaly(M,e) # angle in the orbit focus between the perigee and the satellite direction
    psi = np.arctan2((np.sqrt(1 - e**2))*np.sin(eta), np.cos(eta) - e) # compute psi
    r   = a*( 1 - e*np.cos(eta)) # compute radius r
	
	# Compute the coordinates of the satellite in ORS and store it in coord_ORS
    xORS = r*np.cos(psi)
    yORS = r*np.sin(psi)
    zORS = 0
    coord_ORS[:,Dt] = [xORS, yORS, zORS]

    r_mean[Dt] = r

    # Compute rotation angles omega, i, OMEGA (parameters)
    w = w0 + wdot*(Dt_i - t_start)
    i = i0 + idot*(Dt_i - t_start)
    Omega = Omega0 + (Omegadot - OmegaEdot)*(Dt_i - t_start)

    # Compute the rotation matrices required to transform from ORS to ITRF
    R_Omega = np.array([[np.cos(-Omega), np.sin(-Omega), 0],
                       [-np.sin(-Omega), np.cos(-Omega), 0], 
                       [0, 0, 1]])
    R_i     = np.array([[1, 0, 0],
                       [0, np.cos(-i), -np.sin(-i)],
                       [0, np.sin(-i), np.cos(-i)]])
    R_w     = np.array([[np.cos(-w), np.sin(-w), 0],
                       [-np.sin(-w), np.cos(-w), 0], 
                       [0, 0, 1]])
    
    # Compute the coordinates of the satellites in ITRF and store it in coord_ORS
    coord_ITRF[:,Dt] = R_Omega @ R_i @ R_w @ coord_ORS[:,Dt]


'''
3) Convert satellite's coordinates from global cartesian [X, Y, Z] to geodetic [latitude, longitude, height]
''' 
# Tip: You can use the cartesian to geodetic function that you implemented for the first exercise
coord_geod = np.array([GC2Geodetic(coord_ITRF[:,i]) for i in range(len(t))])
Lat = coord_geod[:,0] 
Lon = coord_geod[:,1] 
h   = coord_geod[:,2]*1e-3 # KM
# Create 3 lists containing the values of Latitude, Longitude and Height
# Compute the average radius
r_av = np.mean(r_mean)
'''
print('X:',[ round(elem, 4) for elem in coord_ORS[0] ])
print('Y:',[ round(elem, 4) for elem in coord_ORS[1] ])
print('Z:',[ round(elem, 4) for elem in coord_ORS[2] ])
print('X:',[ round(elem, 4) for elem in coord_ITRF[0] ])
print('Y:',[ round(elem, 4) for elem in coord_ITRF[1] ])
print('Z:',[ round(elem, 4) for elem in coord_ITRF[2] ])
print('latitude: ',[ round(elem, 4) for elem in deg2sex(coord_geod[0]) ] ) # angles already in degree
print('longitude:',[ round(elem, 4) for elem in deg2sex(coord_geod[1]) ] )
print('h:         ',round(coord_geod[2],4))'''

import sys 
orig_stdout = sys.stdout
f = open('C:/Users/Utente/Documents/Python/Positioning_Lab/LAB02_results.txt','w')
sys.stdout = f
print('ORS coordinates')
print('X:')
for elem in coord_ORS[0]:
    print(round(elem, 4),'\n')
print('Y:')
for elem in coord_ORS[1]:
    print(round(elem, 4),'\n' )
print('Z:')
for elem in coord_ORS[2]:
    print(round(elem, 4),'\n' )       
print('ITRF global cartesian coordinates [m]')
print('X:')
for elem in coord_ITRF[0]:
    print(round(elem, 4),'\n' )
print('Y:')
for elem in coord_ITRF[1]:
    print(round(elem, 4),'\n' )
print('Z:')
for elem in coord_ITRF[2]:
    print(round(elem, 4),'\n' )     
print('Geodetic coordinates phi, lambda, h')
print('latitude: ')
for coors in coord_geod[:,0]:
    print([ round(elem, 4) for elem in deg2sex(-coors) ] ,'\n')
print('longitude: ')
for coors in coord_geod[:,1]:
    print([ round(elem, 4) for elem in deg2sex(coors) ] ,'\n')
print('h: ')
for coors in coord_geod[:,2]:
    print( round(coors, 4) ,'\n')
sys.stdout = orig_stdout
f.close()
'''
4) Plot satellite's daily trajectory with basemap
'''
# REQUIRED LIBRARIES FOR THIS SOLUTION: pandas and geopandas
import pandas as pd
import geopandas as gpd

# Realize a dataframe containing satellite coordinates

#data={'time':t,'Latitude':Lat,'Longitude':Lon}
df = pd.DataFrame()
df['time'] = t
df['Latitude'] = Lat
df['Longitude'] = Lon 

# Transform the DataFrame in a GeoDataFrame
gdf = gpd.GeoDataFrame(df, geometry=gpd.points_from_xy(df.Longitude,-df.Latitude), crs = 3857)

# Load the basemap
world = gpd.read_file('c:/Users/Utente/Documents/Python/Positioning_Lab/LAB02/world/world.shp')

# Plot the trajectory with world basemap
fig, ax = plt.subplots(figsize = (15,15))
world.plot(ax = ax)
gdf.plot(ax = ax, marker='o', color='red')
ax.set(xlabel='Longitude', ylabel='Latitude', title='Satellite daily trajectory')

'''
5) Plot height of the satellite 
'''

fig, ax = plt.subplots(figsize=(10,6))
ax.set(xlabel='seconds in one day (00:00 - 23:59 = 86400 sec)', ylabel='[km]', title='ellipsoidic height variations [km] around mean height = '+str(h.mean())+' [km]')
ax.plot(t, h - h.mean(), '-', color='blue')
plt.show()

'''
6) Print on a text file the results

print('Clock-offset\n', clock_offsets)
print(' ')
print('Satellite coordinates - ORS')
print('X \n',xORS)
print('Y \n',yORS)
print('Z \n',0)
print('Average radius \n',r)
print(' ')
print('Rotation angles (parameters)')
print('w \n',w)
print('i \n',i)
print('Omega \n',Omega)
print(' ')
print('Satellite coordinates - ITRF')
print('X \n',coord_ITRF[:,0])
print('Y \n',coord_ITRF[:,1])
print('Z \n',coord_ITRF[:,2])
print(' ')
print('Satellite daily trajectory - Geodetic Coordinates')
print('Longitude \n',Lon)
print('Latitude \n',Lat)
print('Heigth [km]\n',h)
print(' ')
print('Ellipsoidic height variations [km] around mean height = '+str(mean_h)+' [km] \n',h - h.mean())
print('Seconds in one day (00:00 - 23:59 = 86400 sec) \n',t)

'''

'''
memfile = StringIO()
t = np.arange(0.0, 2.0, 0.01)
s = 1 + np.sin(2*np.pi*t)
plt.plot(t, s)
plt.savefig(memfile)

document = Document()
document.add_heading('Report',0)
document.add_picture(memfile, width=Inches(1.25))

document.save('report.docx')
memfile.close()
'''
